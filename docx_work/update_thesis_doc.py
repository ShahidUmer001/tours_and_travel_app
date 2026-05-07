from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ACCESS_DATE = "30 April 2026"


def set_cell_margins(cell, top=0, start=45, bottom=0, end=45):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def clear_paragraph(paragraph):
    paragraph_element = paragraph._element
    for child in list(paragraph_element):
        paragraph_element.remove(child)


def collapse_cell_paragraphs(cell):
    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    if len(parts) <= 1:
        return

    first = cell.paragraphs[0]
    first_style = first.style
    first_alignment = first.alignment
    clear_paragraph(first)
    first.add_run(" ".join(parts))
    first.style = first_style
    first.alignment = first_alignment

    for paragraph in list(cell.paragraphs[1:]):
        paragraph._element.getparent().remove(paragraph._element)


def remove_empty_rows(table):
    for row in list(table.rows):
        if all(not cell.text.strip() for cell in row.cells):
            row._tr.getparent().remove(row._tr)


def update_reference_dates(document):
    in_references = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "REFERENCES":
            in_references = True
            continue
        if not in_references:
            continue
        if "[Online]. Available:" in text and "[Accessed:" not in text:
            paragraph.add_run(f" [Accessed: {ACCESS_DATE}].")


def find_abbreviation_table(document):
    for table in document.tables:
        if not table.rows or len(table.columns) != 2:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header != ["Abbreviation", "Full Form"]:
            continue
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "PKR" in table_text and "MBaaS" in table_text:
            return table
    raise RuntimeError("Target abbreviation table not found.")


def format_abbreviation_table(table, available_width):
    remove_empty_rows(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    first_col_width = Inches(1.3)
    second_col_width = available_width - first_col_width

    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.width = first_col_width if col_index == 0 else second_col_width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

            if col_index == 1 and row_index > 0:
                collapse_cell_paragraphs(cell)

            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                )

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    input_path = Path(
        r"C:\Users\HP\tours_and_travel_app\docx_work\Final_Year_Project_Tours_Travel_Thesis_APPROVAL_TOC_FIXED 11.docx"
    )
    output_path = Path(
        r"C:\Users\HP\tours_and_travel_app\docx_work\Final_Year_Project_Tours_Travel_Thesis_APPROVAL_TOC_FIXED 11_updated.docx"
    )

    document = Document(input_path)

    update_reference_dates(document)

    section = document.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    format_abbreviation_table(find_abbreviation_table(document), available_width)

    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
