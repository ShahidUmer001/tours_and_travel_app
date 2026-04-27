from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]


def normalize(text: str) -> str:
    return " ".join(text.split()).strip()


def copy_reference_style_parts(reference_docx: Path, output_docx: Path):
    parts_to_copy = {
        "word/styles.xml",
        "word/theme/theme1.xml",
        "word/fontTable.xml",
    }

    with zipfile.ZipFile(output_docx, "r") as source_zip:
        file_map = {name: source_zip.read(name) for name in source_zip.namelist()}

    with zipfile.ZipFile(reference_docx, "r") as ref_zip:
        for part in parts_to_copy:
            if part in ref_zip.namelist():
                file_map[part] = ref_zip.read(part)

    with zipfile.ZipFile(output_docx, "w") as target_zip:
        for name, data in file_map.items():
            target_zip.writestr(name, data)


def set_doc_margins_from_reference(doc: Document, reference: Document):
    ref_section = reference.sections[0]
    for section in doc.sections:
        section.top_margin = ref_section.top_margin
        section.bottom_margin = ref_section.bottom_margin
        section.left_margin = ref_section.left_margin
        section.right_margin = ref_section.right_margin


def set_run_font(paragraph, *, name: str | None = None, size: float | None = None, bold: bool | None = None, italic: bool | None = None):
    for run in paragraph.runs:
        if name is not None:
            run.font.name = name
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic


def set_body_format(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(6)


def set_caption_format(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(4)


def set_heading_format(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)


def style_exists(doc: Document, style_name: str) -> bool:
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


def set_style_if_exists(paragraph, doc: Document, style_name: str):
    if style_exists(doc, style_name):
        paragraph.style = doc.styles[style_name]


def find_first_heading_index(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        text = normalize(paragraph.text)
        if text in {
            "APPROVAL FOR SUBMISSION",
            "DECLARATION OF ORIGINALITY",
            "DECLARATION OF ORGINALITY",
            "ABSTRACT",
        }:
            return idx
    return 0


def configure_cover_page(doc: Document):
    first_heading_index = find_first_heading_index(doc)
    cover = doc.paragraphs[:first_heading_index]

    for paragraph in cover:
        text = normalize(paragraph.text)
        if not text:
            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)

        if text == "Tours & Travel Mobile Application":
            set_style_if_exists(paragraph, doc, "Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, name="Cambria", size=22.5, bold=True)
        elif text.lower().startswith("submitted by"):
            paragraph.text = "Submitted by"
            set_style_if_exists(paragraph, doc, "Body Text")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, size=12)
        elif text in {"Shahid Umer", "Musharaf"}:
            set_style_if_exists(paragraph, doc, "Heading 2")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, size=16, bold=True)
        elif re.match(r"^\(AUIC-", text):
            set_style_if_exists(paragraph, doc, "Heading 1")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, size=16, bold=True)
        elif text.lower().startswith("supervised by"):
            paragraph.text = "Supervised by"
            set_style_if_exists(paragraph, doc, "Body Text")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, size=12)
        elif text == "Dr. Amjad Khan":
            paragraph.text = "Dr. Amjad Khan Assistant Professor"
            set_style_if_exists(paragraph, doc, "Heading 2")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, size=16, bold=True)
        elif text == "Assistant Professor":
            paragraph.text = ""
        elif "final year project report submitted" in text.lower():
            set_style_if_exists(paragraph, doc, "Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, name="Times New Roman", size=12)
        elif text == "Department of Computing":
            paragraph.text = "Department of Computing Abasyn University Islamabad Campus"
            set_style_if_exists(paragraph, doc, "Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, name="Times New Roman", size=12)
        elif text == "Abasyn University Islamabad Campus":
            paragraph.text = ""
        elif text == "April 2026":
            set_style_if_exists(paragraph, doc, "Heading 2")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph, size=16, bold=True)


def screen_explanation(title: str) -> str:
    explanations = {
        "Login Screen": "This screen allows the user to enter email and password and access the app after successful authentication.",
        "Sign Up Screen": "This screen collects the basic user details and creates a new account for the Tours and Travel application.",
        "Social Authentication Options": "This screen area shows alternative sign-in methods like Google, Facebook, and Apple for quick access.",
        "Home Screen": "This screen works as the main dashboard and shows destinations, packages, services, and quick travel options.",
        "Single Destination Cards on Home": "This section highlights featured destinations so the user can open a specific place and view more details.",
        "Multi-City Tour Section": "This section presents tour packages that include more than one city and guide the user toward the booking flow.",
        "Home Services Grid": "This part of the home screen gives direct access to hotel booking, car booking, maps, weather, chat, and other tools.",
        "Bottom Navigation on Home": "The bottom navigation helps the user move between the main areas of the app without losing context.",
        "Destination Detail Screen": "This screen shows the selected destination with travel information, highlights, pricing, and a booking option.",
        "Destination Highlights and Booking Call-to-Action": "This part of the destination page focuses on key travel points and the main action button for booking.",
        "Multi-City Hotel Selection Screen": "This screen lets the user review hotel options linked with the selected tour package before continuing.",
        "Tour Booking Summary Screen": "This screen summarizes the selected package, travel details, and user input before payment or confirmation.",
        "Transport Selection Screen": "This screen allows the user to choose a transport option based on route, price, and travel preference.",
        "Hotel Search Screen": "This screen helps the user search hotels by destination or city and move toward available accommodation options.",
        "Hotel Selection Screen": "This screen shows the hotel list and allows the user to compare and select the most suitable stay option.",
        "Hotel Booking Summary and Price View": "This screen section presents the selected hotel details, room information, and the final price summary.",
        "Car Booking Screen": "This screen supports city-to-city car booking by letting the user enter route and vehicle details.",
        "Car Route and Fare Details": "This section explains the selected route, estimated travel details, and the calculated fare for the booking.",
        "Car Fare Details": "This section explains the selected route, estimated travel details, and the calculated fare for the booking.",
        "Payment Screen Using Card Method": "This screen shows the payment form for card-based payment and lets the user complete the booking securely.",
        "Payment Screen Using Wallet Method": "This screen demonstrates wallet-based payment methods such as JazzCash and EasyPaisa for booking confirmation.",
        "Booking Success Screen": "This screen confirms that the booking has been completed and gives the user a clear success message.",
        "Booking History Main View": "This screen shows the main booking history list where the user can review saved and recent bookings.",
        "Booking History Filtered Tabs": "This part of the booking history allows the user to switch between tours, hotels, and car bookings.",
        "Profile Screen": "This screen displays the user profile and provides access to account details, preferences, and logout actions.",
        "Profile Image Upload or Settings Area": "This section lets the user update the profile image and manage account-related settings inside the app.",
        "Map Screen": "This screen uses map support to help the user view places, tourist locations, and current travel context.",
        "Tourist Spot Detail Bottom Sheet": "This section appears as a bottom sheet and presents more information about a selected tourist spot on the map.",
        "Add Hotel Screen": "This screen allows a user to add a hotel listing with details such as name, location, price, and contact information.",
        "Add Car Screen": "This screen allows a user to add a car listing with route-related, vehicle, and contact details.",
        "My Listings Hotels Tab": "This screen tab shows all hotel listings added by the current user inside the app.",
        "My Listings Cars Tab": "This screen tab shows all car listings added by the current user inside the app.",
        "Example User Hotel Listing": "This example shows how a user-submitted hotel record is stored and displayed in the listing area.",
        "Example User Car Listing": "This example shows how a user-submitted car record is stored and displayed in the listing area.",
        "Firebase Data Initializer Screen": "This utility screen is used to add or initialize demo data so the app can be tested more easily.",
    }

    if title in explanations:
        return explanations[title]

    if "Screen" in title:
        return f"This screen shows the {title.lower()} of the Tours and Travel application and supports the related user task."
    if "Section" in title:
        return f"This section highlights the {title.lower()} and helps the user continue the flow more clearly."
    if "Tab" in title:
        return f"This tab gives a focused view of {title.lower()} so the user can review related records more easily."
    return f"This figure explains the {title.lower()} in the Tours and Travel mobile application."


def insert_chapter5_screen_content(doc: Document):
    paragraphs = list(doc.paragraphs)
    start_idx = None
    end_idx = None

    for idx, paragraph in enumerate(paragraphs):
        text = normalize(paragraph.text)
        if text == "Chapter 5 System Implementation":
            start_idx = idx
        elif start_idx is not None and text == "Chapter 6 Testing, Validation, and Evaluation":
            end_idx = idx
            break

    if start_idx is None or end_idx is None:
        return

    target_paragraphs = []
    for paragraph in paragraphs[start_idx:end_idx]:
        text = normalize(paragraph.text)
        if re.match(r"^Figure 5\.\d+:", text):
            target_paragraphs.append(paragraph)

    for paragraph in reversed(target_paragraphs):
        text = normalize(paragraph.text)
        title = text.split(":", 1)[1].strip() if ":" in text else text
        prev = paragraph._p.getprevious()
        prev_text = ""
        while prev is not None:
            joined = "".join(node.text or "" for node in prev.iter())
            prev_text = normalize(joined)
            if prev_text:
                break
            prev = prev.getprevious()

        if prev_text == title:
            continue

        heading_paragraph = paragraph.insert_paragraph_before(title)
        set_style_if_exists(heading_paragraph, doc, "Heading 4")
        set_heading_format(heading_paragraph)
        set_run_font(heading_paragraph, size=12, bold=True)

        body_paragraph = paragraph.insert_paragraph_before(screen_explanation(title))
        set_style_if_exists(body_paragraph, doc, "Body Text")
        set_body_format(body_paragraph)
        set_run_font(body_paragraph, name="Times New Roman", size=12)


def format_inserted_chapter5_headings(doc: Document):
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        text = normalize(paragraph.text)
        if not re.match(r"^Figure 5\.\d+:", text):
            continue

        found = []
        j = idx - 1
        while j >= 0 and len(found) < 2:
            candidate = normalize(paragraphs[j].text)
            if candidate:
                found.append(paragraphs[j])
            j -= 1

        if len(found) < 2:
            continue

        explanation_para = found[0]
        title_para = found[1]
        title_text = text.split(":", 1)[1].strip() if ":" in text else text

        if normalize(title_para.text) == title_text:
            set_style_if_exists(title_para, doc, "Heading 4")
            set_heading_format(title_para)
            set_run_font(title_para, size=12, bold=True)

            set_style_if_exists(explanation_para, doc, "Body Text")
            set_body_format(explanation_para)
            set_run_font(explanation_para, name="Times New Roman", size=12)


def restyle_paragraphs(doc: Document):
    front_titles_h1 = {
        "APPROVAL FOR SUBMISSION",
        "DECLARATION OF ORIGINALITY",
        "DECLARATION OF ORGINALITY",
        "ABSTRACT",
    }
    front_titles_h2 = {
        "SUBMISSION AND COPYRIGHT",
        "DEDICATION",
        "ACKNOWLEDGEMENT",
        "References",
        "REFERENCES",
    }
    toc_titles = {
        "TABLE OF CONTENTS",
        "LIST OF ABBREVIATIONS",
        "LIST OF TABLES",
        "LIST OF FIGURES",
    }

    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if not text:
            continue

        if text in front_titles_h1:
            set_style_if_exists(paragraph, doc, "Heading 1")
            set_heading_format(paragraph)
            set_run_font(paragraph, size=16, bold=True)
        elif text in front_titles_h2:
            if text == "REFERENCES":
                paragraph.text = "References"
            set_style_if_exists(paragraph, doc, "Heading 2")
            set_heading_format(paragraph)
            set_run_font(paragraph, size=16, bold=True)
        elif text in toc_titles:
            set_style_if_exists(paragraph, doc, "Heading 3")
            set_heading_format(paragraph)
            set_run_font(paragraph, name="Arial", size=14, bold=True)
        elif text.startswith("Chapter "):
            set_style_if_exists(paragraph, doc, "Heading 2")
            set_heading_format(paragraph)
            set_run_font(paragraph, size=16, bold=True)
        elif re.match(r"^\d+\.\d+", text):
            set_style_if_exists(paragraph, doc, "Heading 4")
            set_heading_format(paragraph)
            set_run_font(paragraph, size=12, bold=True)
        elif text.startswith("Figure "):
            set_style_if_exists(paragraph, doc, "Heading 4")
            set_caption_format(paragraph)
            set_run_font(paragraph, size=12, bold=True)
        elif text.startswith("Table "):
            set_style_if_exists(paragraph, doc, "Heading 4")
            set_caption_format(paragraph)
            set_run_font(paragraph, size=12, bold=True)
        elif paragraph.style and paragraph.style.name == "List Paragraph":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(2)
        else:
            set_style_if_exists(paragraph, doc, "Body Text")
            set_body_format(paragraph)
            set_run_font(paragraph, name="Times New Roman", size=12)


def restyle_tables(doc: Document):
    for table in doc.tables:
        try:
            table.style = "Normal Table"
        except Exception:
            pass
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if row_idx == 0:
                        set_style_if_exists(paragraph, doc, "Heading 4")
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_run_font(paragraph, size=11, bold=True)
                    else:
                        set_style_if_exists(paragraph, doc, "Body Text")
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        set_run_font(paragraph, name="Times New Roman", size=11)


def transform_document(source_docx: Path, reference_docx: Path, output_docx: Path):
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source_docx, output_docx)

    doc = Document(str(output_docx))
    reference = Document(str(reference_docx))

    set_doc_margins_from_reference(doc, reference)
    insert_chapter5_screen_content(doc)
    restyle_paragraphs(doc)
    format_inserted_chapter5_headings(doc)
    configure_cover_page(doc)
    restyle_tables(doc)

    doc.save(str(output_docx))
    copy_reference_style_parts(reference_docx, output_docx)


def main():
    source_docx = Path(r"C:\Users\HP\Desktop\Shahid_Umer_And_Musharaf_Tours_Travel_Thesis.docx")
    reference_docx = Path(r"C:\Users\HP\Downloads\Expedition Management System changing (3).docx")
    output_docx = ROOT / "docs" / "Shahid_Umer_And_Musharaf_Tours_Travel_Thesis_ExpeditionStyle.docx"
    transform_document(source_docx, reference_docx, output_docx)
    print(f"Created: {output_docx}")


if __name__ == "__main__":
    main()
