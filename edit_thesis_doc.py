from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph


INPUT_PATH = Path(r"C:\Users\HP\tours_and_travel_app\Final_Year_Project_Tours_&_Travel_Thesis_SaveTest 1_working.docx")
OUTPUT_PATH = Path(r"C:\Users\HP\tours_and_travel_app\Final_Year_Project_Tours_&_Travel_Thesis_Edited.docx")

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


ABSTRACT_TEXT = (
    "Pakistan offers a wide range of tourism opportunities, including natural landscapes, cultural sites, "
    "historical landmarks, and emerging domestic travel routes, yet visitors often depend on disconnected "
    "services for destination discovery, hotel reservations, transport selection, and booking follow-up. "
    "This fragmentation increases planning effort, reduces convenience, and weakens the overall travel "
    "experience. To address this problem, the Tours & Travel Mobile Application was developed as an "
    "integrated academic prototype that brings major tourism-related functions into a single mobile platform. "
    "The application is built with Flutter and Dart and uses Firebase services where available for "
    "authentication, cloud storage, profile management, and data persistence. It allows users to explore "
    "destinations, review tour packages, reserve hotels, book city-to-city transport, view location support "
    "through maps, manage booking history, update personal profiles, and create local hotel or car listings. "
    "The system follows a modular structure that separates user interface elements, services, models, and "
    "utility components to improve maintainability, readability, and future enhancement potential. A hybrid "
    "demonstration approach is also adopted so the project can continue to function in limited testing "
    "environments through local storage and fallback session handling when some cloud features are unavailable. "
    "The completed application demonstrates how cross-platform mobile development can support practical tourism "
    "digitization within the Pakistani context while remaining suitable for academic presentation, evaluation, "
    "and extension. It also emphasizes user-friendly interface design and a service-based "
    "architecture that helps individual modules be explained, tested, and improved independently. Functional "
    "validation of the prototype shows that the application can support realistic travel planning scenarios in a "
    "single workflow, even when some production-grade integrations are intentionally simplified for final year "
    "project constraints. The project therefore serves both as a usable demonstration system and as a structured "
    "foundation for future work such as live payment integration, recommendation features, admin-level controls, "
    "analytics, stronger backend rules, and broader deployment support."
)


TOC_ENTRIES = [
    ("APPROVAL FOR SUBMISSION", "ii"),
    ("DECLARATION OF ORIGINALITY", "iv"),
    ("SUBMISSION AND COPYRIGHTS", "v"),
    ("ACKNOWLEDGMENT", "vi"),
    ("ABSTRACT", "vii"),
    ("TABLE OF CONTENTS", "viii"),
    ("LIST OF ABBREVIATIONS", "ix"),
    ("LIST OF TABLES", "xii"),
    ("LIST OF FIGURES", "xvi"),
    ("CHAPTER 1 INTRODUCTION", "16"),
    ("CHAPTER 2 LITERATURE REVIEW", "20"),
    ("CHAPTER 3 REQUIREMENT ENGINEERING AND FEASIBILITY STUDY", "23"),
    ("CHAPTER 4 SYSTEM ANALYSIS AND DESIGN", "27"),
    ("CHAPTER 5 SYSTEM IMPLEMENTATION", "37"),
    ("CHAPTER 6 TESTING STRATEGY", "75"),
    ("CHAPTER 7 RESULTS, LIMITATIONS, AND DEPLOYMENT CONSIDERATIONS", "79"),
    ("CHAPTER 8 CONCLUSION AND FUTURE ENHANCEMENTS", "82"),
    ("REFERENCES", "84"),
]


FIGURE_TARGETS = [
    (
        "Figure 4.2: Use Case Diagram",
        "The Use Case perspective is critical because it moves analysis from simply screens to user intents. "
        "Instead of discussing the system only in terms of interface components it questions what actions are "
        "valuable to the user and what ordered steps the system has to support.",
    ),
    (
        "Figure 4.3: Data Flow Diagram Level 0",
        "The Use Case perspective is critical because it moves analysis from simply screens to user intents. "
        "Instead of discussing the system only in terms of interface components it questions what actions are "
        "valuable to the user and what ordered steps the system has to support.",
    ),
    (
        "Figure 4.4: Data Flow Diagram Level 1",
        "The Use Case perspective is critical because it moves analysis from simply screens to user intents. "
        "Instead of discussing the system only in terms of interface components it questions what actions are "
        "valuable to the user and what ordered steps the system has to support.",
    ),
    (
        "Figure 4.11: Navigation Structure Diagram",
        "Navigation design is screen based and imperative. On startup the application leads to authentication or "
        "directly to the homepage upon authentication. The bottom navigation and card-based action flows allow "
        "access to tours, history, profile, map and listings from the homepage. The importance of navigation "
        "design in the context of a travel application, users flip-flop between finding information and "
        "confirming selection, cannot be understated. In this case, I have focused on observable actions, "
        "familiar navigations and summary screens to avoid confusion. Every booking flow seeks to maintain "
        "context, allowing the user to recognize what it is that he has chosen and why the system requires the "
        "next input.",
    ),
    (
        "Figure 4.10: Firestore Schema Diagram",
        "The schema is kept deliberately simplistic for academic manageability whilst still encompassing the "
        "relevant relationship: A user could relate to profile information, bookings relate to certain "
        "destination or choices of service, or listings related to the owner ID. Local storage is used in the "
        "specific places where the need to manage a set of user-created listing or where a fallback mechanism "
        "should exist in the absence of the cloud.",
    ),
    (
        "Figure 4.9: Class Diagram",
        "This service-based approach towards designing classes gives the codebase a tangible form that students "
        "and reviewers can easily relate to. The clear separation between data representations and behaviors "
        "helps users understand that models simply contain data (a booking, a hotel) while services handle "
        "behavior associated with them (retrieval, modification, display etc.).",
    ),
    (
        "Figure 4.5: Activity Diagram for Login",
        "Because destinations themselves are aspirational goods, emotional design plays a vital role in "
        "applications oriented towards traveling. Hence the UI is not overlooked but image and motion are "
        "strategically used to set the tone and deliver ambiance, all the while maintaining clear interaction "
        "rules. This can be observed in the homepage, the destination cards as well as the booking confirmation "
        "process.",
    ),
    (
        "Figure 4.6: Activity Diagram for Tour Booking",
        "Because destinations themselves are aspirational goods, emotional design plays a vital role in "
        "applications oriented towards traveling. Hence the UI is not overlooked but image and motion are "
        "strategically used to set the tone and deliver ambiance, all the while maintaining clear interaction "
        "rules. This can be observed in the homepage, the destination cards as well as the booking confirmation "
        "process.",
    ),
    (
        "Figure 4.7: Sequence Diagram for Authentication",
        "Even if this is not a product release the design of validation logic is of primary concern as it "
        "defines user-trust and prevents ill-fated state changes. The following components form part of the "
        "present design: client-side validation, awareness of the local session, and a restricted mode of "
        "storage-operations. It will be possible to extend this design to a secure format by adding "
        "authorization, encryption, and server-side checks.",
    ),
    (
        "Figure 4.8: Sequence Diagram for Payment and Confirmation",
        "Even if this is not a product release the design of validation logic is of primary concern as it "
        "defines user-trust and prevents ill-fated state changes. The following components form part of the "
        "present design: client-side validation, awareness of the local session, and a restricted mode of "
        "storage-operations. It will be possible to extend this design to a secure format by adding "
        "authorization, encryption, and server-side checks.",
    ),
    (
        "Figure 4.12: Deployment Diagram",
        "Even if this is not a product release the design of validation logic is of primary concern as it "
        "defines user-trust and prevents ill-fated state changes. The following components form part of the "
        "present design: client-side validation, awareness of the local session, and a restricted mode of "
        "storage-operations. It will be possible to extend this design to a secure format by adding "
        "authorization, encryption, and server-side checks.",
    ),
]


def paragraph_text(p: CT_P) -> str:
    return "".join(p.xpath(".//w:t/text()")).strip()


def has_drawing(p: CT_P) -> bool:
    return bool(p.xpath(".//w:drawing") or p.xpath(".//w:pict"))


def body_paragraphs(doc: Document) -> list[CT_P]:
    return [el for el in doc.element.body.iterchildren() if isinstance(el, CT_P)]


def find_body_paragraph(doc: Document, text: str, *, after_text: str | None = None) -> CT_P:
    seen_after = after_text is None
    for p in body_paragraphs(doc):
        current = paragraph_text(p)
        if not seen_after and current == after_text:
            seen_after = True
        elif seen_after and current == text:
            return p
    raise ValueError(f"Could not find paragraph: {text!r}")


def clone_blank_paragraph(template: CT_P, style_name: str = "Normal") -> CT_P:
    new_p = deepcopy(template)
    ppr = new_p.find(qn("w:pPr"))
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        new_p.insert(0, ppr)
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        pstyle = OxmlElement("w:pStyle")
        ppr.insert(0, pstyle)
    pstyle.set(qn("w:val"), style_name)
    return new_p


def set_simple_entry_text(p: CT_P, title: str, page: str) -> None:
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)

    run = OxmlElement("w:r")
    text1 = OxmlElement("w:t")
    text1.text = title
    run.append(text1)

    tab_run = OxmlElement("w:r")
    tab = OxmlElement("w:tab")
    tab_run.append(tab)

    page_run = OxmlElement("w:r")
    text2 = OxmlElement("w:t")
    text2.text = page
    page_run.append(text2)

    p.append(run)
    p.append(tab_run)
    p.append(page_run)


def insert_after(anchor: CT_P, new_elements: list[CT_P]) -> CT_P:
    parent = anchor.getparent()
    idx = parent.index(anchor)
    for offset, element in enumerate(new_elements, start=1):
        parent.insert(idx + offset, element)
    return new_elements[-1]


def move_after(anchor: CT_P, elements: list[CT_P]) -> CT_P:
    parent = anchor.getparent()
    for element in elements:
        if element.getparent() is not None:
            element.getparent().remove(element)
    idx = parent.index(anchor)
    for offset, element in enumerate(elements, start=1):
        parent.insert(idx + offset, element)
    return elements[-1]


def next_nonempty_after(doc: Document, heading_text: str) -> CT_P:
    seen_heading = False
    for p in body_paragraphs(doc):
        current = paragraph_text(p)
        if current == heading_text:
            seen_heading = True
            continue
        if seen_heading and current:
            return p
    raise ValueError(f"No content found after heading: {heading_text}")


def previous_nonempty(elements: list[CT_P], start_idx: int, stop_idx: int) -> int | None:
    for i in range(start_idx, stop_idx, -1):
        if paragraph_text(elements[i]):
            return i
    return None


def next_nonempty(elements: list[CT_P], start_idx: int, stop_idx: int) -> int | None:
    for i in range(start_idx, stop_idx):
        if paragraph_text(elements[i]):
            return i
    return None


def extract_figure_block(doc: Document, caption_text: str, cluster_end_text: str) -> list[CT_P]:
    elems = body_paragraphs(doc)
    caption_idx = next(i for i, p in enumerate(elems) if paragraph_text(p) == caption_text)
    cluster_end_idx = next(i for i, p in enumerate(elems) if paragraph_text(p) == cluster_end_text)

    prev_nonempty_idx = previous_nonempty(elems, caption_idx - 1, -1)
    search_start = (prev_nonempty_idx or -1) + 1

    image_idx = None
    for i in range(search_start, caption_idx):
        if has_drawing(elems[i]):
            image_idx = i
            break
    if image_idx is None:
        raise ValueError(f"No image paragraph found for {caption_text}")

    next_nonempty_idx = next_nonempty(elems, caption_idx + 1, cluster_end_idx + 1)
    end_idx = next_nonempty_idx if next_nonempty_idx is not None else cluster_end_idx

    block = elems[image_idx:end_idx]
    kept: list[CT_P] = []
    blank_count = 0
    for el in block:
        txt = paragraph_text(el)
        if has_drawing(el) or txt == caption_text:
            kept.append(el)
            blank_count = 0
            continue
        if not txt and blank_count < 1:
            kept.append(el)
            blank_count += 1
    if not any(paragraph_text(el) == caption_text for el in kept):
        kept.append(elems[caption_idx])
    return kept


def clean_cluster(doc: Document, start_after_text: str, end_before_text: str) -> None:
    start = find_body_paragraph(doc, start_after_text)
    end = find_body_paragraph(doc, end_before_text)
    parent = doc.element.body
    started = False
    for el in list(parent.iterchildren()):
        if el is start:
            started = True
            continue
        if el is end:
            break
        if started and isinstance(el, CT_P):
            txt = paragraph_text(el)
            if not txt and not has_drawing(el):
                parent.remove(el)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def ensure_word_count(text: str) -> None:
    count = len(text.split())
    if not 280 <= count <= 300:
        raise ValueError(f"Abstract word count is {count}, expected 280-300.")


def update_abstract(doc: Document) -> None:
    ensure_word_count(ABSTRACT_TEXT)
    set_paragraph_text(doc.paragraphs[102], ABSTRACT_TEXT)
    for idx in (101, 103, 104, 105):
        doc.paragraphs[idx].text = ""


def update_tables(doc: Document) -> None:
    tables = doc.tables

    # Approval tables
    tables[0].cell(0, 0).text = "Dr. Amjad Khan"
    tables[0].cell(3, 0).text = "Pending signature"

    tables[1].cell(0, 0).text = "To be assigned by department"
    tables[1].cell(3, 0).text = "Pending signature"

    tables[2].cell(0, 0).text = "To be assigned by department"
    tables[2].cell(3, 0).text = "Pending signature"

    # Split/continued tables
    tables[6].cell(6, 1).text = "Map and tourist marker exploration"

    tables[7].cell(0, 0).text = "Requirement ID"
    tables[7].cell(0, 1).text = "Description"
    tables[7].cell(0, 2).text = "Primary Module"

    tables[10].cell(3, 1).text = "Cloud features may fail during demo"
    tables[10].cell(3, 2).text = "Local persistence and seed data approaches"

    tables[11].cell(0, 0).text = "Risk"
    tables[11].cell(0, 1).text = "Impact"
    tables[11].cell(0, 2).text = "Mitigation"


def fix_headings(doc: Document) -> None:
    replacements = {
        "DECLARATION OF ORGINALITY": "DECLARATION OF ORIGINALITY",
        "Chapter3": "Chapter 3",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in replacements:
            para.text = replacements[text]


def update_toc(doc: Document) -> None:
    anchor = find_body_paragraph(doc, "TABLE OF CONTENTS")
    p_end = find_body_paragraph(doc, "LIST OF ABBREVIATIONS", after_text="TABLE OF CONTENTS")
    parent = doc.element.body

    p_start = anchor.getnext()
    if p_start is None:
        raise ValueError("No paragraph found after TABLE OF CONTENTS.")
    template = deepcopy(p_start)

    for el in list(parent.iterchildren()):
        if el is anchor:
            continue
        if el is p_end:
            break
        if el.getparent() is parent and parent.index(el) > parent.index(anchor):
            parent.remove(el)

    first_template = deepcopy(template)
    entries: list[CT_P] = []
    for title, page in TOC_ENTRIES:
        entry = deepcopy(first_template)
        set_simple_entry_text(entry, title, page)
        entries.append(entry)

    for entry in reversed(entries):
        anchor.addnext(entry)


def update_list_titles(doc: Document) -> None:
    for para in doc.paragraphs:
        if para.text.strip() == "LIST OF ABBREVIATION":
            para.text = "LIST OF ABBREVIATIONS"
        elif para.text.strip() == "LIST OF TABLE":
            para.text = "LIST OF TABLES"


def reorganize_chapter_four_figures(doc: Document) -> None:
    cluster_end_text = "Chapter 5 System Implementation"
    blocks = {
        caption: extract_figure_block(doc, caption, cluster_end_text)
        for caption, _ in FIGURE_TARGETS
    }

    blank_template = next(p for p in body_paragraphs(doc) if not paragraph_text(p) and not has_drawing(p))
    last_inserted: dict[str, CT_P] = {}

    for caption, anchor_text in FIGURE_TARGETS:
        anchor = last_inserted.get(anchor_text)
        if anchor is None:
            anchor = find_body_paragraph(doc, anchor_text)
        moved = move_after(anchor, blocks[caption])
        spacer = clone_blank_paragraph(blank_template, "BodyText")
        moved = insert_after(moved, [spacer])
        last_inserted[anchor_text] = moved

    clean_cluster(doc, "This stage builds on the requirements and maps them to an actual system architecture. It dictates user interaction within the application, data location, service interaction and diagrams can be used to define the software architecture. Decisions made in this stage influence the next development chapter.", "Chapter 5 System Implementation")


def main() -> None:
    doc = Document(str(INPUT_PATH))

    fix_headings(doc)
    update_abstract(doc)
    update_tables(doc)
    update_list_titles(doc)
    reorganize_chapter_four_figures(doc)
    update_toc(doc)

    doc.save(str(OUTPUT_PATH))
    print(f"Saved edited thesis to: {OUTPUT_PATH}")
    print(f"Abstract word count: {len(ABSTRACT_TEXT.split())}")


if __name__ == "__main__":
    main()
