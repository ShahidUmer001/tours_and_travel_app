from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "docs" / "Shahid_Source_Copy.docx"
OUTPUT_DOCX = ROOT / "docs" / "Shahid_Umer_And_Musharaf_Tours_Travel_Thesis_ExpeditionExact_v2.docx"


def normalize(text: str) -> str:
    return " ".join(text.split()).strip()


def insert_paragraph_before(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def set_paragraph_text(paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_body_paragraph_style(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def set_heading4_style(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        run.bold = True
        run.font.size = Pt(12)


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if normalize(paragraph.text) == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def same_paragraph(left, right) -> bool:
    return left._element is right._element


def section_paragraphs_between(doc: Document, start_text: str, end_text: str) -> list[str]:
    start = False
    items: list[str] = []
    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text == start_text:
            start = True
            continue
        if start and text == end_text:
            break
        if start and text:
            items.append(text)
    return items


def extract_source_data(source_doc: Document):
    approval_items = section_paragraphs_between(source_doc, "APPROVAL FOR SUBMISSION", "DECLARATION OF ORIGINALITY")
    declaration_items = section_paragraphs_between(source_doc, "DECLARATION OF ORIGINALITY", "SUBMISSION AND COPYRIGHT")
    submission_items = section_paragraphs_between(source_doc, "SUBMISSION AND COPYRIGHT", "DEDICATION")
    acknowledgement_items = section_paragraphs_between(source_doc, "ACKNOWLEDGEMENT", "ABSTRACT")
    abstract_items = section_paragraphs_between(source_doc, "ABSTRACT", "TABLE OF CONTENTS")
    tables_items = section_paragraphs_between(source_doc, "LIST OF TABLES", "LIST OF FIGURES")
    figures_items = section_paragraphs_between(source_doc, "LIST OF FIGURES", "Chapter 1 Introduction")

    abbreviations_rows = []
    for idx, paragraph in enumerate(source_doc.paragraphs):
        if normalize(paragraph.text) == "LIST OF ABBREVIATIONS":
            if idx + 1 < len(source_doc.tables) + 200:
                break

    # Table 3 in the source is the abbreviations table.
    abbr_table = source_doc.tables[2]
    for row in abbr_table.rows:
        abbreviations_rows.append([normalize(cell.text) for cell in row.cells])

    return {
        "approval": approval_items,
        "declaration": declaration_items,
        "submission": submission_items,
        "acknowledgement": acknowledgement_items,
        "abstract": abstract_items,
        "tables": tables_items,
        "figures": figures_items,
        "abbreviations": abbreviations_rows,
    }


def replace_cover_page(doc: Document):
    replacements = {
        "Expedition Management System": "Tours & Travel Mobile Application",
        "Hissam Umair": "Shahid Umer",
        "(AUIC-FL20-BSSE-4953)": "(AUIC-22SG-BSCS-6732)",
        "Mubashir Hussain": "Musharaf",
        "(AUIC-FL20-BSSE-4979)": "(AUIC-22SG-BSCS-7029)",
        "Mr. Abdul Hannan Assistant Professor": "Dr. Amjad Khan Assistant Professor",
        "A final year project report submitted in partial fulfillment of the requirement for the award of Bachelor of Science in Software Engineering": "A final year project report submitted in partial fulfillment of the requirement for the award of Bachelor of Science in Computer Science",
        "October 2024": "April 2026",
    }
    for paragraph in doc.paragraphs:
        text = normalize(paragraph.text)
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])


def replace_between_with_body_paragraphs(
    doc: Document,
    start_heading: str,
    end_heading: str,
    new_items: list[str],
    *,
    style: str = "Body Text",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
):
    start_para = find_paragraph(doc, start_heading)
    end_para = find_paragraph(doc, end_heading)

    to_remove = []
    collect = False
    for paragraph in doc.paragraphs:
        if same_paragraph(paragraph, start_para):
            collect = True
            continue
        if same_paragraph(paragraph, end_para):
            break
        if collect:
            if normalize(paragraph.text):
                to_remove.append(paragraph)

    for paragraph in to_remove:
        remove_paragraph(paragraph)

    anchor = end_para
    for text in new_items:
        para = insert_paragraph_before(anchor, text, style)
        if style == "Body Text":
            set_body_paragraph_style(para, align=align)
        else:
            para.alignment = align


def replace_declaration_section(doc: Document, source_data):
    start = find_paragraph(doc, "DECLARATION OF ORGINALITY")
    end = find_paragraph(doc, "Submission and Copyrights")

    to_remove = []
    collect = False
    for paragraph in doc.paragraphs:
        if same_paragraph(paragraph, start):
            collect = True
            continue
        if same_paragraph(paragraph, end):
            break
        if collect and normalize(paragraph.text):
            to_remove.append(paragraph)

    for paragraph in to_remove:
        remove_paragraph(paragraph)

    anchor = end
    items = [
        source_data["declaration"][0],
        source_data["declaration"][1],
        "Signature:",
        "Shahid Umer",
        "Signature:",
        "Musharaf",
        "Date: ________________________________",
    ]
    for text in items:
        para = insert_paragraph_before(anchor, text, "Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def replace_submission_section(doc: Document, source_data):
    start = find_paragraph(doc, "Submission and Copyrights")
    end = find_paragraph(doc, "Acknowledgment")

    to_remove = []
    collect = False
    for paragraph in doc.paragraphs:
        if same_paragraph(paragraph, start):
            collect = True
            continue
        if same_paragraph(paragraph, end):
            break
        if collect and normalize(paragraph.text):
            to_remove.append(paragraph)

    for paragraph in to_remove:
        remove_paragraph(paragraph)

    anchor = end
    items = [
        source_data["submission"][0] if source_data["submission"] else "This project report is submitted to the Department of Computing at Abasyn University Islamabad Campus as partial fulfillment of the requirements for the award of the undergraduate degree of Bachelor of Science in Computer Science.",
        "Shahid Umer (AUIC-22SG-BSCS-6732)",
        "Musharaf (AUIC-22SG-BSCS-7029)",
        "Dr. Amjad Khan (Supervisor)",
        "Assistant Professor, Department of Computing",
        "Copyright © 2026 Shahid Umer and Musharaf. All rights reserved.",
    ]
    for idx, text in enumerate(items):
        para = insert_paragraph_before(anchor, text, "Normal" if idx < 4 else "Body Text")
        if idx == 0:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif idx in {1, 2, 3}:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            set_body_paragraph_style(para, align=WD_ALIGN_PARAGRAPH.CENTER)


def replace_acknowledgement_and_abstract(doc: Document, source_data):
    replace_between_with_body_paragraphs(
        doc,
        "Acknowledgment",
        "ABSTRACT",
        source_data["acknowledgement"],
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    replace_between_with_body_paragraphs(
        doc,
        "ABSTRACT",
        "TABLE OF CONTENTS",
        source_data["abstract"],
        style="Body Text",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def replace_approval_section(doc: Document, source_data):
    committee_para = find_paragraph(doc, "Approval Committee:")
    start = find_paragraph(doc, "APPROVAL FOR SUBMISSION")

    to_remove = []
    collect = False
    for paragraph in doc.paragraphs:
        if same_paragraph(paragraph, start):
            collect = True
            continue
        if same_paragraph(paragraph, committee_para):
            break
        if collect and normalize(paragraph.text):
            to_remove.append(paragraph)

    for paragraph in to_remove:
        remove_paragraph(paragraph)

    anchor = committee_para
    body_items = [text for text in source_data["approval"] if text != "Approval Committee"]
    for text in body_items:
        para = insert_paragraph_before(anchor, text, "Normal")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def replace_tables_section_items(doc: Document, start_heading: str, end_heading: str, items: list[str]):
    start = find_paragraph(doc, start_heading)
    end = find_paragraph(doc, end_heading)
    to_remove = []
    collect = False
    for paragraph in doc.paragraphs:
        if same_paragraph(paragraph, start):
            collect = True
            continue
        if same_paragraph(paragraph, end):
            break
        if collect and normalize(paragraph.text):
            to_remove.append(paragraph)

    for paragraph in to_remove:
        remove_paragraph(paragraph)

    anchor = end
    for text in items:
        para = insert_paragraph_before(anchor, text, "Body Text")
        set_body_paragraph_style(para, align=WD_ALIGN_PARAGRAPH.LEFT)


def replace_abbreviations_table(doc: Document, rows: list[list[str]]):
    heading = find_paragraph(doc, "LIST OF ABBREVIATIONS")
    table_index = None
    for idx, table in enumerate(doc.tables):
        # find first table that appears after the heading in document order by scanning elements
        if table._element.getprevious() is not None:
            prev_text = normalize("".join(table._element.getprevious().itertext()))
            if "LIST OF ABBREVIATIONS" in prev_text:
                table_index = idx
                break
    if table_index is None:
        table_index = 4

    table = doc.tables[table_index]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    header = rows[0]
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = header[i] if i < len(header) else ""

    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if i < len(cells):
                cells[i].text = value

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if row_idx == 0:
                    paragraph.style = doc.styles["Heading 4"]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.style = doc.styles["Body Text"]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def update_front_tables(doc: Document):
    replacements = {
        "Hissam Umair": "Shahid Umer",
        "Mubashir Hussain": "Musharaf",
        "Mr. Abdul Hannan Assistant Professor": "Dr. Amjad Khan Assistant Professor",
        "Mr. Abdul Hannan (Supervisor)": "Dr. Amjad Khan (Supervisor)",
        "(AUIC-FL20-BSSE-4953)": "(AUIC-22SG-BSCS-6732)",
        "(AUIC-FL20-BSSE-4979)": "(AUIC-22SG-BSCS-7029)",
        "AUIC-FL20-BSSE-4953": "AUIC-22SG-BSCS-6732",
        "AUIC-FL20-BSSE-4979": "AUIC-22SG-BSCS-7029",
        "Software Engineering": "Computer Science",
    }
    for table in doc.tables[:5]:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for old, new in replacements.items():
                    text = text.replace(old, new)
                cell.text = text


def main():
    source_doc = Document(str(SOURCE_DOCX))
    output_doc = Document(str(OUTPUT_DOCX))

    source_data = extract_source_data(source_doc)

    replace_cover_page(output_doc)
    replace_approval_section(output_doc, source_data)
    replace_declaration_section(output_doc, source_data)
    replace_submission_section(output_doc, source_data)
    replace_acknowledgement_and_abstract(output_doc, source_data)
    replace_tables_section_items(output_doc, "LIST OF TABLES", "LIST OF FIGURES", source_data["tables"])
    replace_tables_section_items(output_doc, "LIST OF FIGURES", "Chapter 1 Introduction", source_data["figures"])
    replace_abbreviations_table(output_doc, source_data["abbreviations"])
    update_front_tables(output_doc)

    output_doc.save(str(OUTPUT_DOCX))
    print(f"Updated: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
